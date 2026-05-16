"""
gerar_vistoria_entrada_saida.py
────────────────────────────────────────────────────────────────────────────
Gera o DOCX da vistoria com colunas ENTREGA e DEVOLUÇÃO no mesmo documento.

Diferente do gerar_vistoria_nova original (que mexia no XML cru), esta
versão usa python-docx puro, porque o template é controlado por nós
(VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx, gerado por criar_template.py).
É mais simples, mais legível e menos propenso a quebrar.

Como usar (a partir do app.py):

    from gerar_vistoria_entrada_saida import gerar_vistoria_entrada_saida

    dados = {
        # ── Dados fixos do contrato ─────────────────────────────────────
        "contrato_id":       "0042",
        "cliente_nome":      "JOÃO DA SILVA",
        "cliente_telefone":  "(83) 99999-0000",
        "cliente_endereco":  "Rua X, 123 — João Pessoa/PB",
        "preenchido_por":    "Ana Karolina",
        "veiculo":           "VW Gol Special",
        "placa":             "ELY-4D83",
        "cor":               "BRANCO",
        "ano":               "2022/2023",
        "chassi":            "9BW...",
        "numero_motor":      "...",

        # ── Bloco de ENTREGA ────────────────────────────────────────────
        "data_entrada":         "10/05/2026 14:30",
        "hodometro_entrada":    "32.450",
        "combustivel_entrada":  "3/4",
        "obs_entrada":          "Veículo limpo, sem riscos.",
        "sintomas_entrada":     "",
        "responsavel_entrada":  "Ana Karolina",
        "acessorios_entrada":   {"acc_calotas": "S", "acc_buzina": "S", ...},

        # ── Fotos: dicionário ângulo → caminho  ─────────────────────────
        # As chaves seguem ANGULOS_FOTO (ver abaixo). Envie apenas os
        # ângulos que o operador fotografou — os ausentes ficam em branco.
        # Formato alternativo legado: lista simples de caminhos, ainda
        # aceita e é convertida automaticamente para dict sem legenda.
        "fotos_entrada": {
            "frontal":          "/path/frente.jpg",
            "traseira":         "/path/traseira.jpg",
            "lateral_dir":      "/path/lat_dir.jpg",
            "lateral_esq":      "/path/lat_esq.jpg",
            "painel":           "/path/painel.jpg",
            "hodometro":        "/path/hodometro.jpg",
            "estepe":           "/path/estepe.jpg",
            "teto":             "/path/teto.jpg",
        },

        # ── Bloco de DEVOLUÇÃO (pode vir vazio na 1ª geração) ───────────
        "data_saida":           "20/05/2026 11:00",
        "hodometro_saida":      "33.120",
        "combustivel_saida":    "1/4",
        "obs_saida":            "Para-choque dianteiro arranhado.",
        "sintomas_saida":       "AC fazendo barulho ao ligar.",
        "responsavel_saida":    "Ana Karolina",
        "acessorios_saida":     {...},
        "fotos_saida":          {...},
    }

    resumo = gerar_vistoria_entrada_saida(
        dados,
        caminho_saida="contratos/VISTORIA_ELY4D83.docx",
        template_path="docx_templates/VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx",
    )
    # resumo == {
    #   "arquivo": "...", "status": "completa"|"pendente_saida",
    #   "divergencias": [(label, entrada, saida, motivo), ...]
    # }

Status do documento:
- "pendente_saida"  → só a entrega foi preenchida (cliente acabou de retirar)
- "completa"        → entrega + devolução

Status de cada acessório (coluna "Status / Observação"):
- igual              → cinza   "OK — sem alteração"
- S→N (sumiu)        → vermelho "Item ausente na devolução"
- S→A (avariou)      → vermelho "Avariado na devolução"
- N→A (surgiu avaria)→ vermelho "Surgiu avaria"
- N→S, A→S           → verde   "Reposto/recuperado"
- só entrada preench → cinza-claro "Aguardando devolução"

Fotos — padronização por ângulo:
- Cada foto é inserida em um slot fixo de FOTO_LARGURA_CM × FOTO_ALTURA_CM.
- A imagem é recortada/redimensionada pelo script para caber exatamente
  nesse slot, preservando proporção (letter-box com fundo branco).
- As fotos são exibidas em grade 2 colunas, com legenda embaixo.
- Slots sem foto recebem um retângulo cinza "— sem foto —".
- A ordem dos ângulos segue ANGULOS_FOTO (pode ser personalizada).
"""
from __future__ import annotations

import io
import shutil
from pathlib import Path
from typing import Optional, Union

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Padronização de fotos ────────────────────────────────────────────────────
# Dimensões do slot de cada foto no documento (em cm).
# Todas as fotos são redimensionadas para caber exatamente nesse espaço,
# com letter-box branco caso a proporção não bata.
FOTO_LARGURA_CM: float = 8.0   # largura do slot no DOCX
FOTO_ALTURA_CM:  float = 6.0   # altura do slot no DOCX  (proporção 4:3)

# Número de colunas na grade de fotos.
FOTO_COLUNAS: int = 2

# Ângulos reconhecidos e suas legendas, na ordem de exibição.
# Adicione/remova entradas conforme o fluxo do seu app.
ANGULOS_FOTO: list[tuple[str, str]] = [
    ("frontal",       "Frontal"),
    ("traseira",      "Traseira"),
    ("lateral_dir",   "Lateral Direita"),
    ("lateral_esq",   "Lateral Esquerda"),
    ("painel",        "Painel / Interior"),
    ("hodometro",     "Hodômetro"),
    ("estepe",        "Estepe"),
    ("teto",          "Teto / Vidro"),
    ("motor",         "Motor"),
    ("mala",          "Porta-malas"),
    ("dano_1",        "Dano 1"),
    ("dano_2",        "Dano 2"),
]

# Cor de fundo do slot vazio (cinza claro) — RGB
_CINZA_SLOT = (220, 220, 220)


# ── Lista padrão dos 30 acessórios (na mesma ordem do template antigo) ──────
ACESSORIOS = [
    "acc_calotas",          "acc_buzina",           "acc_doc_crlv",
    "acc_triangulo",        "acc_antena",           "acc_sensor_re",
    "acc_som",              "acc_tapetes",          "acc_limpadores",
    "acc_chave_roda",       "acc_vidros_eletricos", "acc_oleo_motor",
    "acc_alarme",           "acc_lampadas",         "acc_macaco",
    "acc_estepe",           "acc_gnv",              "acc_agua",
    "acc_borr_psg_dir",     "acc_borr_mtr_dir",     "acc_asa_dd",
    "acc_asa_td",           "acc_tapete_mala",      "acc_tampa_parachoque",
    "acc_borr_psg_tras",    "acc_borr_mtr_tras",    "acc_asa_de",
    "acc_asa_te",           "acc_bagagito",         "acc_lingueta",
]

# Cores
COR_VERDE       = RGBColor(0x1A, 0x6B, 0x1A)
COR_VERMELHO    = RGBColor(0xCC, 0x00, 0x00)
COR_AMARELO     = RGBColor(0xD4, 0xAC, 0x0D)
COR_CINZA       = RGBColor(0x55, 0x55, 0x55)
COR_CINZA_CLARO = RGBColor(0x99, 0x99, 0x99)

PLACEHOLDER_VAZIO = "—"


# ─────────────────────────────────────────────────────────────────────────────
# Lógica de comparação entrada × saída
# ─────────────────────────────────────────────────────────────────────────────
def _cor_valor(v: str) -> RGBColor:
    v = (v or "").upper().strip()
    return {"S": COR_VERDE, "N": COR_VERMELHO, "A": COR_AMARELO}.get(v, COR_CINZA_CLARO)


def _calcular_status_acessorio(entrada: str, saida: str) -> tuple[str, RGBColor]:
    e = (entrada or "").upper().strip()
    s = (saida   or "").upper().strip()

    if not e and not s:
        return "", COR_CINZA_CLARO
    if e and not s:
        return "Aguardando devolução", COR_CINZA_CLARO
    if s and not e:
        return "Sem registro de entrega", COR_AMARELO
    if e == s:
        return "OK — sem alteração", COR_CINZA

    if e == "S" and s == "N":
        return "Item ausente na devolução", COR_VERMELHO
    if e == "S" and s == "A":
        return "Avariado na devolução", COR_VERMELHO
    if e == "N" and s == "A":
        return "Surgiu avaria", COR_VERMELHO
    if (e, s) in (("N", "S"), ("A", "S"), ("A", "N")):
        return "Reposto/recuperado", COR_VERDE
    return f"Mudou: {e} → {s}", COR_AMARELO


# ─────────────────────────────────────────────────────────────────────────────
# Substituição de placeholders em python-docx
# ─────────────────────────────────────────────────────────────────────────────
def _set_celula(cell, texto: str, *, cor: Optional[RGBColor] = None,
                bold: Optional[bool] = None, size: Optional[int] = None,
                align: Optional[str] = None):
    paragraphs = cell.paragraphs
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = paragraphs[0]
    if align:
        p.alignment = {
            'left':   WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(align, p.alignment)

    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    run = p.add_run(texto if texto else PLACEHOLDER_VAZIO)
    if cor is not None:
        run.font.color.rgb = cor
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


def _substituir_placeholder_no_celula(cell, mapa: dict, *, cor_padrao=COR_CINZA,
                                       size_padrao=10):
    texto_completo = cell.text
    placeholders_encontrados = [(ph, val, cor, sz) for ph, (val, cor, sz) in mapa.items()
                                 if f"[{ph}]" in texto_completo]
    if not placeholders_encontrados:
        return

    if len(placeholders_encontrados) == 1:
        ph, val, cor, sz = placeholders_encontrados[0]
        bare = texto_completo.strip() == f"[{ph}]"
        if bare:
            _set_celula(cell, val if val else PLACEHOLDER_VAZIO,
                        cor=cor, bold=True, size=sz)
            return

    novo_texto = texto_completo
    for ph, val, _cor, _sz in placeholders_encontrados:
        novo_texto = novo_texto.replace(f"[{ph}]", val if val else PLACEHOLDER_VAZIO)
    paragraphs = cell.paragraphs
    if len(paragraphs) == 1:
        p = paragraphs[0]
        antigo_align = p.alignment
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        linhas = novo_texto.split('\n')
        for i, linha in enumerate(linhas):
            run = p.add_run(linha)
            run.font.color.rgb = cor_padrao
            run.font.size = Pt(size_padrao)
            if i < len(linhas) - 1:
                run.add_break()
        if antigo_align is not None:
            p.alignment = antigo_align
    else:
        partes = novo_texto.split('\n')
        for i, p in enumerate(paragraphs):
            txt = partes[i] if i < len(partes) else ""
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run(txt)
            run.font.color.rgb = cor_padrao
            run.font.size = Pt(size_padrao)


def _substituir_em_paragrafo(par, mapa: dict, *, cor=COR_CINZA, size=10):
    if not par.text:
        return
    novo = par.text
    mudou = False
    for ph, (val, _cor, _sz) in mapa.items():
        token = f"[{ph}]"
        if token in novo:
            novo = novo.replace(token, val if val else PLACEHOLDER_VAZIO)
            mudou = True
    if mudou:
        for r in list(par.runs):
            r._element.getparent().remove(r._element)
        run = par.add_run(novo)
        run.font.color.rgb = cor
        run.font.size = Pt(size)


# ─────────────────────────────────────────────────────────────────────────────
# Normalização de imagens — Pillow (opcional mas recomendado)
# ─────────────────────────────────────────────────────────────────────────────
def _normalizar_foto(caminho: str) -> io.BytesIO:
    """
    Abre a imagem, corrige orientação EXIF, redimensiona para caber no slot
    FOTO_LARGURA_CM × FOTO_ALTURA_CM (em 96 dpi) com letter-box branco, e
    devolve um BytesIO com o JPEG resultante.

    Se Pillow não estiver instalado, devolve a imagem original sem alterar.
    """
    try:
        from PIL import Image, ImageOps
    except ImportError:
        buf = io.BytesIO(Path(caminho).read_bytes())
        buf.seek(0)
        return buf

    DPI = 96
    alvo_w = int(round(FOTO_LARGURA_CM / 2.54 * DPI))
    alvo_h = int(round(FOTO_ALTURA_CM  / 2.54 * DPI))

    img = Image.open(caminho)
    img = ImageOps.exif_transpose(img)
    img = img.convert("RGB")

    img.thumbnail((alvo_w, alvo_h), Image.LANCZOS)

    canvas = Image.new("RGB", (alvo_w, alvo_h), (255, 255, 255))
    off_x = (alvo_w - img.width)  // 2
    off_y = (alvo_h - img.height) // 2
    canvas.paste(img, (off_x, off_y))

    buf = io.BytesIO()
    canvas.save(buf, format="JPEG", quality=85, optimize=True)
    buf.seek(0)
    return buf


def _slot_vazio() -> Optional[io.BytesIO]:
    """
    Gera um retângulo cinza claro para representar um ângulo sem foto.
    Requer Pillow; devolve None se não disponível.
    """
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        return None

    DPI = 96
    w = int(round(FOTO_LARGURA_CM / 2.54 * DPI))
    h = int(round(FOTO_ALTURA_CM  / 2.54 * DPI))

    img = Image.new("RGB", (w, h), _CINZA_SLOT)
    draw = ImageDraw.Draw(img)
    texto = "— sem foto —"
    bbox = draw.textbbox((0, 0), texto)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((w - tw) // 2, (h - th) // 2), texto, fill=(150, 150, 150))

    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────────────────
# Inserção de fotos padronizadas em grade
# ─────────────────────────────────────────────────────────────────────────────
def _normalizar_entrada_fotos(fotos) -> dict[str, str]:
    """
    Aceita dois formatos de entrada:
    • dict  {"frontal": "/path/a.jpg", ...}  → filtra caminhos válidos
    • list  ["/path/a.jpg", ...]             → mapeia na ordem de ANGULOS_FOTO
    Retorna sempre {angulo: caminho}.
    """
    if not fotos:
        return {}
    if isinstance(fotos, dict):
        return {k: v for k, v in fotos.items() if v and Path(v).exists()}
    chaves = [k for k, _ in ANGULOS_FOTO]
    return {
        (chaves[i] if i < len(chaves) else f"extra_{i}"): caminho
        for i, caminho in enumerate(fotos)
        if caminho and Path(caminho).exists()
    }


def _set_tcW(tc, twips: int) -> None:
    """Define a largura da célula sem gerar elemento <w:tcW> duplicado."""
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _inserir_fotos_no_marcador(doc: Document, marcador: str, fotos) -> None:
    """
    Localiza o parágrafo `marcador` e o substitui por uma grade de fotos
    padronizadas (FOTO_COLUNAS colunas, FOTO_LARGURA_CM × FOTO_ALTURA_CM por slot).

    Sempre exibe todos os 12 ângulos de ANGULOS_FOTO:
    • slots com foto → imagem normalizada (letter-box branco)
    • slots sem foto → retângulo cinza "— sem foto —"
    Ângulos extras (fora de ANGULOS_FOTO) são acrescentados ao final.
    """
    fotos_dict = _normalizar_entrada_fotos(fotos)

    alvo = None
    for p in doc.paragraphs:
        if p.text.strip() == marcador:
            alvo = p
            break
    if alvo is None:
        return

    for r in list(alvo.runs):
        r._element.getparent().remove(r._element)

    # ── Monta lista com TODOS os ângulos padrão + extras ─────────────────────
    # caminho=None → slot vazio (placeholder cinza)
    slots: list[tuple[str, str, Optional[str]]] = [
        (chave, legenda, fotos_dict.get(chave))
        for chave, legenda in ANGULOS_FOTO
    ]
    chaves_padrao = {k for k, _ in ANGULOS_FOTO}
    for chave, caminho in fotos_dict.items():
        if chave not in chaves_padrao:
            slots.append((chave, chave.replace("_", " ").title(), caminho))

    # ── Cria tabela de fotos ──────────────────────────────────────────────────
    n_cols      = FOTO_COLUNAS
    n_rows      = -(-len(slots) // n_cols)
    col_w_twips = int(FOTO_LARGURA_CM * 567)

    tabela = doc.add_table(rows=n_rows * 2, cols=n_cols)

    tblPr = tabela._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for lado in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{lado}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    for i, (chave, legenda, caminho) in enumerate(slots):
        linha_foto    = (i // n_cols) * 2
        linha_legenda = linha_foto + 1
        col           = i % n_cols

        cell_foto    = tabela.cell(linha_foto, col)
        cell_legenda = tabela.cell(linha_legenda, col)

        _set_tcW(cell_foto._tc,    col_w_twips)
        _set_tcW(cell_legenda._tc, col_w_twips)

        p_foto = cell_foto.paragraphs[0]
        p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_foto = p_foto.add_run()
        try:
            if caminho:
                buf = _normalizar_foto(caminho)
            else:
                buf = _slot_vazio()
            if buf is not None:
                run_foto.add_picture(buf, width=Cm(FOTO_LARGURA_CM))
            else:
                run_foto.text = "(sem foto)"
                run_foto.font.size = Pt(8)
                run_foto.font.color.rgb = COR_CINZA_CLARO
        except Exception as exc:
            run_foto.text = f"(erro: {exc})"
            run_foto.font.size = Pt(8)
            run_foto.font.color.rgb = COR_VERMELHO

        p_leg = cell_legenda.paragraphs[0]
        p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_leg = p_leg.add_run(legenda)
        run_leg.bold = True
        run_leg.font.size = Pt(8)
        run_leg.font.color.rgb = COR_CINZA

    parent      = alvo._element.getparent()
    tbl_element = tabela._tbl
    parent.remove(tbl_element)
    alvo._element.addnext(tbl_element)


# ─────────────────────────────────────────────────────────────────────────────
# Função principal
# ─────────────────────────────────────────────────────────────────────────────
def gerar_vistoria_entrada_saida(
    dados: dict,
    caminho_saida: str,
    template_path: Optional[str] = None,
) -> dict:
    """
    Gera o DOCX a partir do template e devolve um resumo.
    """
    template = Path(template_path or
                    Path(__file__).parent / "docx_templates" /
                    "VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx")
    if not template.exists():
        raise FileNotFoundError(f"Template não encontrado: {template}")

    acc_in  = dados.get("acessorios_entrada") or {}
    acc_out = dados.get("acessorios_saida")   or {}

    tem_saida = any([
        dados.get("data_saida"), dados.get("hodometro_saida"),
        dados.get("combustivel_saida"), acc_out,
        dados.get("obs_saida"), dados.get("sintomas_saida"),
        dados.get("fotos_saida"),
    ])
    status_doc = "completa" if tem_saida else "pendente_saida"

    mapa: dict[str, tuple[str, RGBColor, int]] = {}

    campos_simples = [
        "contrato_id", "cliente_nome", "cliente_telefone",
        "cliente_endereco", "preenchido_por",
        "veiculo", "placa", "cor", "ano", "chassi", "numero_motor",
        "data_entrada", "data_saida",
        "hodometro_entrada", "hodometro_saida",
        "combustivel_entrada", "combustivel_saida",
        "obs_entrada", "obs_saida",
        "sintomas_entrada", "sintomas_saida",
        "responsavel_entrada", "responsavel_saida",
    ]
    for campo in campos_simples:
        mapa[campo] = (str(dados.get(campo) or ""), COR_CINZA, 10)

    divergencias = []
    for key in ACESSORIOS:
        e = acc_in.get(key, "")
        s = acc_out.get(key, "")
        status_txt, status_cor = _calcular_status_acessorio(e, s)
        mapa[f"{key}_entrada"] = (e, _cor_valor(e), 9)
        mapa[f"{key}_saida"]   = (s, _cor_valor(s), 9)
        mapa[f"{key}_status"]  = (status_txt, status_cor, 9)

        if status_txt and "OK" not in status_txt and "Aguardando" not in status_txt:
            label = key.replace("acc_", "").replace("_", " ").title()
            divergencias.append((label, e or "—", s or "—", status_txt))

    doc = Document(str(template))

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _substituir_placeholder_no_celula(cell, mapa)

    for par in doc.paragraphs:
        _substituir_em_paragrafo(par, mapa)

    _inserir_fotos_no_marcador(doc, "[FOTOS_ENTRADA]",
                                dados.get("fotos_entrada", []))
    _inserir_fotos_no_marcador(doc, "[FOTOS_SAIDA]",
                                dados.get("fotos_saida", []))

    saida = Path(caminho_saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(saida))

    return {
        "arquivo":      str(saida),
        "status":       status_doc,
        "divergencias": divergencias,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Conversão para PDF (mesma abordagem do gerar_contrato.py)
# ─────────────────────────────────────────────────────────────────────────────
def docx_para_pdf(caminho_docx: str) -> Optional[str]:
    """
    Converte o DOCX para PDF usando docx2pdf (Windows + Word) ou
    libreoffice (Linux). Retorna o caminho do PDF gerado, ou None.
    """
    import platform, subprocess
    caminho_pdf = str(Path(caminho_docx).with_suffix(".pdf"))
    sistema = platform.system().lower()
    try:
        if sistema == "windows":
            from docx2pdf import convert
            convert(caminho_docx, caminho_pdf)
        else:
            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", str(Path(caminho_docx).parent),
                caminho_docx
            ], check=True, timeout=60)
        return caminho_pdf if Path(caminho_pdf).exists() else None
    except Exception as e:
        print(f"[gerar_vistoria_entrada_saida] PDF falhou: {e}")
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Execução direta — gera um exemplo
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    exemplo = {
        "contrato_id":       "0042",
        "cliente_nome":      "JOÃO DA SILVA SANTOS",
        "cliente_telefone":  "(83) 99999-0000",
        "cliente_endereco":  "Rua das Acácias, 123 — Bessa, João Pessoa/PB",
        "preenchido_por":    "Ana Karolina",
        "veiculo":           "VW Gol Special 1.0",
        "placa":             "ELY-4D83",
        "cor":               "BRANCO",
        "ano":               "2022/2023",
        "chassi":            "9BWZZZ377VT004251",
        "numero_motor":      "CHZ123456",

        "data_entrada":         "10/05/2026 14:30",
        "hodometro_entrada":    "32.450",
        "combustivel_entrada":  "3/4",
        "obs_entrada":          "Veículo limpo, sem riscos visíveis.",
        "sintomas_entrada":     "",
        "responsavel_entrada":  "Ana Karolina",
        "acessorios_entrada":   {k: "S" for k in ACESSORIOS},

        "data_saida":         "20/05/2026 11:00",
        "hodometro_saida":    "33.120",
        "combustivel_saida":  "1/4",
        "obs_saida":          "Para-choque dianteiro arranhado no canto esquerdo.",
        "sintomas_saida":     "Ar-condicionado fazendo barulho ao ligar.",
        "responsavel_saida":  "Ana Karolina",
        "acessorios_saida": {
            **{k: "S" for k in ACESSORIOS},
            "acc_calotas":   "N",
            "acc_tapetes":   "A",
            "acc_estepe":    "N",
            "acc_triangulo": "N",
        },
        "fotos_entrada": {
            # "frontal":     "/path/frente.jpg",
            # "traseira":    "/path/traseira.jpg",
            # "painel":      "/path/painel.jpg",
            # "hodometro":   "/path/hodometro.jpg",
        },
        "fotos_saida": {
            # "frontal":     "/path/frente_dev.jpg",
            # "dano_1":      "/path/amassado.jpg",
        },
    }
    saida = sys.argv[1] if len(sys.argv) > 1 else "VISTORIA_EXEMPLO.docx"
    resumo = gerar_vistoria_entrada_saida(
        exemplo,
        caminho_saida=saida,
        template_path="VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx",
    )
    print(f"Gerado: {resumo['arquivo']}")
    print(f"Status: {resumo['status']}")
    print(f"Divergências encontradas: {len(resumo['divergencias'])}")
    for label, e, s, motivo in resumo['divergencias']:
        print(f"  • {label}: {e} → {s}  ({motivo})")
