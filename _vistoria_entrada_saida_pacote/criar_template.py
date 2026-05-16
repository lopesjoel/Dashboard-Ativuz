"""
Cria o template DOCX da vistoria com colunas Entrada e Saída.
Roda uma vez para gerar VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Lista padrão de 30 acessórios (mantendo nomenclatura do template antigo) ──
ACESSORIOS = [
    ("acc_calotas",          "Calotas"),
    ("acc_buzina",           "Buzina"),
    ("acc_doc_crlv",         "DOC. CRLV"),
    ("acc_triangulo",        "Triângulo Sinaliz."),
    ("acc_antena",           "Antena"),
    ("acc_sensor_re",        "Sensor de Ré"),
    ("acc_som",              "Som / Alto-falante"),
    ("acc_tapetes",          "Tapetes"),
    ("acc_limpadores",       "Limpadores"),
    ("acc_chave_roda",       "Chave de Roda"),
    ("acc_vidros_eletricos", "Vidros Elétricos"),
    ("acc_oleo_motor",       "Óleo do Motor"),
    ("acc_alarme",           "Alarme / Travas"),
    ("acc_lampadas",         "Lâmpadas"),
    ("acc_macaco",           "Macaco Mecânico"),
    ("acc_estepe",           "Estepe"),
    ("acc_gnv",              "Func. GNV"),
    ("acc_agua",             "Água"),
    ("acc_borr_psg_dir",     "Borracha PSG Dir."),
    ("acc_borr_mtr_dir",     "Borracha MTR Dir."),
    ("acc_asa_dd",           "Asa Urubu D.D."),
    ("acc_asa_td",           "Asa Urubu T.D."),
    ("acc_tapete_mala",      "Tapete de Mala"),
    ("acc_tampa_parachoque", "Tampa Para-choque"),
    ("acc_borr_psg_tras",    "Borracha PSG Tras."),
    ("acc_borr_mtr_tras",    "Borracha MTR Tras."),
    ("acc_asa_de",           "Asa Urubu D.E."),
    ("acc_asa_te",           "Asa Urubu T.E."),
    ("acc_bagagito",         "Bagagito"),
    ("acc_lingueta",         "Lingueta"),
]

# ─────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def cell(table, r, c, text, *, bold=False, size=10, bg=None, align='left'):
    tcell = table.cell(r, c)
    tcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tcell.text = ''
    p = tcell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if bg:
        set_cell_bg(tcell, bg)
    return tcell

# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Margens
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)

# ── Cabeçalho ────────────────────────────────────────────────────────────────
h1 = doc.add_paragraph()
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h1.add_run("ANEXO I — CONTRATO DE LOCAÇÃO")
r.bold = True; r.font.size = Pt(12)

h2 = doc.add_paragraph()
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h2.add_run("VISTORIA DE ENTREGA E DEVOLUÇÃO DO VEÍCULO")
r.bold = True; r.font.size = Pt(13)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Contrato [contrato_id]   •   Documento único — preenchido em duas etapas")
r.italic = True; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # espaço

# ── 1. Dados do cliente ──────────────────────────────────────────────────────
p = doc.add_paragraph(); r = p.add_run("1. DADOS DO CLIENTE")
r.bold = True; r.font.size = Pt(11)

t = doc.add_table(rows=3, cols=4); add_borders(t)
cell(t, 0, 0, "Cliente",        bold=True, bg='F2F2F2')
cell(t, 0, 1, "[cliente_nome]")
cell(t, 0, 2, "Telefone",       bold=True, bg='F2F2F2')
cell(t, 0, 3, "[cliente_telefone]")
cell(t, 1, 0, "Endereço",       bold=True, bg='F2F2F2')
t.cell(1, 1).merge(t.cell(1, 3))
cell(t, 1, 1, "[cliente_endereco]")
cell(t, 2, 0, "Preenchido por", bold=True, bg='F2F2F2')
cell(t, 2, 1, "[preenchido_por]")
cell(t, 2, 2, "Contrato",       bold=True, bg='F2F2F2')
cell(t, 2, 3, "[contrato_id]")

doc.add_paragraph()

# ── 2. Dados do veículo ──────────────────────────────────────────────────────
p = doc.add_paragraph(); r = p.add_run("2. DADOS DO VEÍCULO")
r.bold = True; r.font.size = Pt(11)

t = doc.add_table(rows=3, cols=4); add_borders(t)
cell(t, 0, 0, "Veículo",      bold=True, bg='F2F2F2')
cell(t, 0, 1, "[veiculo]")
cell(t, 0, 2, "Placa",        bold=True, bg='F2F2F2')
cell(t, 0, 3, "[placa]")
cell(t, 1, 0, "Cor",          bold=True, bg='F2F2F2')
cell(t, 1, 1, "[cor]")
cell(t, 1, 2, "Ano",          bold=True, bg='F2F2F2')
cell(t, 1, 3, "[ano]")
cell(t, 2, 0, "Chassi",       bold=True, bg='F2F2F2')
cell(t, 2, 1, "[chassi]")
cell(t, 2, 2, "Número Motor", bold=True, bg='F2F2F2')
cell(t, 2, 3, "[numero_motor]")

doc.add_paragraph()

# ── 3. Comparativo Entrada × Saída (medidores) ───────────────────────────────
p = doc.add_paragraph(); r = p.add_run("3. ENTREGA × DEVOLUÇÃO")
r.bold = True; r.font.size = Pt(11)

t = doc.add_table(rows=4, cols=3); add_borders(t)
# Cabeçalho
cell(t, 0, 0, "Item",      bold=True, bg='1F3A5F', align='center')
cell(t, 0, 1, "Entrega",   bold=True, bg='1F3A5F', align='center')
cell(t, 0, 2, "Devolução", bold=True, bg='1F3A5F', align='center')
# colorir texto do cabeçalho de branco
for c in range(3):
    for run in t.cell(0, c).paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

cell(t, 1, 0, "Data e hora", bold=True, bg='F2F2F2')
cell(t, 1, 1, "[data_entrada]",   align='center')
cell(t, 1, 2, "[data_saida]",     align='center')

cell(t, 2, 0, "Hodômetro (km)",   bold=True, bg='F2F2F2')
cell(t, 2, 1, "[hodometro_entrada]", align='center')
cell(t, 2, 2, "[hodometro_saida]",   align='center')

cell(t, 3, 0, "Combustível",      bold=True, bg='F2F2F2')
cell(t, 3, 1, "[combustivel_entrada]", align='center')
cell(t, 3, 2, "[combustivel_saida]",   align='center')

doc.add_paragraph()

# ── 4. Acessórios e equipamentos ─────────────────────────────────────────────
p = doc.add_paragraph(); r = p.add_run("4. ACESSÓRIOS E EQUIPAMENTOS")
r.bold = True; r.font.size = Pt(11)

leg = doc.add_paragraph()
r = leg.add_run("S = Sim, existente     |     N = Não existente     |     A = Avariado")
r.font.size = Pt(8); r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Tabela com 4 colunas: Item | Entrega | Devolução | Status
t = doc.add_table(rows=len(ACESSORIOS) + 1, cols=4); add_borders(t)
# largura aproximada
for col, w in zip(t.columns, [Cm(5.5), Cm(2.5), Cm(2.5), Cm(6.0)]):
    for c in col.cells:
        c.width = w

# Cabeçalho
cell(t, 0, 0, "Acessório",  bold=True, bg='1F3A5F', align='left')
cell(t, 0, 1, "Entrega",    bold=True, bg='1F3A5F', align='center')
cell(t, 0, 2, "Devolução",  bold=True, bg='1F3A5F', align='center')
cell(t, 0, 3, "Status / Observação", bold=True, bg='1F3A5F', align='left')
for c in range(4):
    for run in t.cell(0, c).paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for i, (key, label) in enumerate(ACESSORIOS):
    row = i + 1
    bg = 'FAFAFA' if i % 2 == 0 else None
    cell(t, row, 0, label, size=9, bg=bg)
    cell(t, row, 1, f"[{key}_entrada]", size=9, align='center', bg=bg, bold=True)
    cell(t, row, 2, f"[{key}_saida]",   size=9, align='center', bg=bg, bold=True)
    cell(t, row, 3, f"[{key}_status]",  size=9, bg=bg)

doc.add_paragraph()

# ── 5. Observações ───────────────────────────────────────────────────────────
p = doc.add_paragraph(); r = p.add_run("5. OBSERVAÇÕES")
r.bold = True; r.font.size = Pt(11)

t = doc.add_table(rows=2, cols=2); add_borders(t)
cell(t, 0, 0, "Observações — Entrega",   bold=True, bg='F2F2F2', align='center')
cell(t, 0, 1, "Observações — Devolução", bold=True, bg='F2F2F2', align='center')
cell(t, 1, 0, "[obs_entrada]")
cell(t, 1, 1, "[obs_saida]")

doc.add_paragraph()

# ── 6. Sintomas / Danos ──────────────────────────────────────────────────────
p = doc.add_paragraph(); r = p.add_run("6. SINTOMAS, DANOS OU AVARIAS RELATADOS")
r.bold = True; r.font.size = Pt(11)

t = doc.add_table(rows=2, cols=2); add_borders(t)
cell(t, 0, 0, "Na entrega",   bold=True, bg='F2F2F2', align='center')
cell(t, 0, 1, "Na devolução", bold=True, bg='F2F2F2', align='center')
cell(t, 1, 0, "[sintomas_entrada]")
cell(t, 1, 1, "[sintomas_saida]")

doc.add_paragraph()

# ── 7. Assinaturas ───────────────────────────────────────────────────────────
p = doc.add_paragraph(); r = p.add_run("7. ASSINATURAS")
r.bold = True; r.font.size = Pt(11)

t = doc.add_table(rows=3, cols=2); add_borders(t)
cell(t, 0, 0, "ENTREGA",   bold=True, bg='1F3A5F', align='center')
cell(t, 0, 1, "DEVOLUÇÃO", bold=True, bg='1F3A5F', align='center')
for c in range(2):
    for run in t.cell(0, c).paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
cell(t, 1, 0, "\n\n\n______________________________________\nAssinatura do Cliente\n[cliente_nome]", align='center', size=9)
cell(t, 1, 1, "\n\n\n______________________________________\nAssinatura do Cliente\n[cliente_nome]", align='center', size=9)
cell(t, 2, 0, "\n\n\n______________________________________\nResponsável Ativuz\n[responsavel_entrada]", align='center', size=9)
cell(t, 2, 1, "\n\n\n______________________________________\nResponsável Ativuz\n[responsavel_saida]",  align='center', size=9)

doc.add_paragraph()

# ── 8. Fotos ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph(); r = p.add_run("8. REGISTRO FOTOGRÁFICO")
r.bold = True; r.font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run("FOTOS DA ENTREGA")
r.bold = True; r.font.size = Pt(10)
doc.add_paragraph("[FOTOS_ENTRADA]")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("FOTOS DA DEVOLUÇÃO")
r.bold = True; r.font.size = Pt(10)
doc.add_paragraph("[FOTOS_SAIDA]")

# Salvar
out = "VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx"
doc.save(out)
print(f"Template criado: {out}")
