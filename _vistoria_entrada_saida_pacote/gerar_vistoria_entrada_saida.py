"""
gerar_vistoria_entrada_saida.py
────────────────────────────────────────────────────────────────────────────
Gera o DOCX da vistoria com colunas ENTREGA e DEVOLUÇÃO no mesmo documento.

Diferente do gerar_vistoria_nova original (que mexia no XML cru), esta
versão usa python-docx puro, porque o template é controlado por nós
(VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx, gerado por criar_template.py).
É mais simples, mais legível e menos propenso a quebrar.

Como usar (a partir do app.py):

    from gerar_vistoria_entrada_saida import gerar_vistoria_entrada_saida

    dados = {
        # ── Dados fixos do contrato ─────────────────────────────────────
        "contrato_id":       "0042",
        "cliente_nome":      "JOÃO DA SILVA",
        "cliente_telefone":  "(83) 99999-0000",
        "cliente_endereco":  "Rua X, 123 — João Pessoa/PB",
        "preenchido_por":    "Ana Karolina",
        "veiculo":           "VW Gol Special",
        "placa":             "ELY-4D83",
        "cor":               "BRANCO",
        "ano":               "2022/2023",
        "chassi":            "9BW...",
        "numero_motor":      "...",

        # ── Bloco de ENTREGA ────────────────────────────────────────────
        "data_entrada":         "10/05/2026 14:30",
        "hodometro_entrada":    "32.450",
        "combustivel_entrada":  "3/4",
        "obs_entrada":          "Veículo limpo, sem riscos.",
        "sintomas_entrada":     "",
        "responsavel_entrada":  "Ana Karolina",
        "acessorios_entrada":   {"acc_calotas": "S", "acc_buzina": "S", ...},
        "fotos_entrada":        ["/caminho/foto1.jpg", ...],

        # ── Bloco de DEVOLUÇÃO (pode vir vazio na 1ª geração) ───────────
        "data_saida":           "20/05/2026 11:00",
        "hodometro_saida":      "33.120",
        "combustivel_saida":    "1/4",
        "obs_saida":            "Para-choque dianteiro arranhado.",
        "sintomas_saida":       "AC fazendo barulho ao ligar.",
        "responsavel_saida":    "Ana Karolina",
        "acessorios_saida":     {...},
        "fotos_saida":          [...],
    }

    resumo = gerar_vistoria_entrada_saida(
        dados,
        caminho_saida="contratos/VISTORIA_ELY4D83.docx",
        template_path="docx_templates/VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx",
    )
    # resumo == {
    #   "arquivo": "...", "status": "completa"|"pendente_saida",
    #   "divergencias": [(label, entrada, saida, motivo), ...]
    # }

Status do documento:
- "pendente_saida"  → só a entrega foi preenchida (cliente acabou de retirar)
- "completa"        → entrega + devolução

Status de cada acessório (coluna "Status / Observação"):
- igual              → cinza   "OK — sem alteração"
- S→N (sumiu)        → vermelho "Item ausente na devolução"
- S→A (avariou)      → vermelho "Avariado na devolução"
- N→A (surgiu avaria)→ vermelho "Surgiu avaria"
- N→S, A→S           → verde   "Reposto/recuperado"
- só entrada preench → cinza-claro "Aguardando devolução"
"""
from __future__ import annotations

import shutil
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Lista padrão dos 30 acessórios (na mesma ordem do template antigo) ──────
ACESSORIOS = [
    "acc_calotas",          "acc_buzina",           "acc_doc_crlv",
    "acc_triangulo",        "acc_antena",           "acc_sensor_re",
    "acc_som",              "acc_tapetes",          "acc_limpadores",
    "acc_chave_roda",       "acc_vidros_eletricos", "acc_oleo_motor",
    "acc_alarme",           "acc_lampadas",         "acc_macaco",
    "acc_estepe",           "acc_gnv",              "acc_agua",
    "acc_borr_psg_dir",     "acc_borr_mtr_dir",     "acc_asa_dd",
    "acc_asa_td",           "acc_tapete_mala",      "acc_tampa_parachoque",
    "acc_borr_psg_tras",    "acc_borr_mtr_tras",    "acc_asa_de",
    "acc_asa_te",           "acc_bagagito",         "acc_lingueta",
]

# Cores
COR_VERDE       = RGBColor(0x1A, 0x6B, 0x1A)
COR_VERMELHO    = RGBColor(0xCC, 0x00, 0x00)
COR_AMARELO     = RGBColor(0xD4, 0xAC, 0x0D)
COR_CINZA       = RGBColor(0x55, 0x55, 0x55)
COR_CINZA_CLARO = RGBColor(0x99, 0x99, 0x99)

PLACEHOLDER_VAZIO = "—"


# ─────────────────────────────────────────────────────────────────────────────
# Lógica de comparação entrada × saída
# ─────────────────────────────────────────────────────────────────────────────
def _cor_valor(v: str) -> RGBColor:
    v = (v or "").upper().strip()
    return {"S": COR_VERDE, "N": COR_VERMELHO, "A": COR_AMARELO}.get(v, COR_CINZA_CLARO)


def _calcular_status_acessorio(entrada: str, saida: str) -> tuple[str, RGBColor]:
    e = (entrada or "").upper().strip()
    s = (saida   or "").upper().strip()

    if not e and not s:
        return "", COR_CINZA_CLARO
    if e and not s:
        return "Aguardando devolução", COR_CINZA_CLARO
    if s and not e:
        return "Sem registro de entrega", COR_AMARELO
    if e == s:
        return "OK — sem alteração", COR_CINZA

    if e == "S" and s == "N":
        return "Item ausente na devolução", COR_VERMELHO
    if e == "S" and s == "A":
        return "Avariado na devolução", COR_VERMELHO
    if e == "N" and s == "A":
        return "Surgiu avaria", COR_VERMELHO
    if (e, s) in (("N", "S"), ("A", "S"), ("A", "N")):
        return "Reposto/recuperado", COR_VERDE
    return f"Mudou: {e} → {s}", COR_AMARELO


# ─────────────────────────────────────────────────────────────────────────────
# Substituição de placeholders em python-docx
# ─────────────────────────────────────────────────────────────────────────────
def _set_celula(cell, texto: str, *, cor: Optional[RGBColor] = None,
                bold: Optional[bool] = None, size: Optional[int] = None,
                align: Optional[str] = None):
    """
    Reescreve o conteúdo da célula preservando o parágrafo (para manter
    fundo, bordas, alinhamento). Mantém só o primeiro parágrafo.
    """
    # Limpa todos os parágrafos, mantém o primeiro vazio
    paragraphs = cell.paragraphs
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = paragraphs[0]
    if align:
        p.alignment = {
            'left':   WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(align, p.alignment)

    # Remove runs antigos
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    run = p.add_run(texto if texto else PLACEHOLDER_VAZIO)
    if cor is not None:
        run.font.color.rgb = cor
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = Pt(size)


def _substituir_placeholder_no_celula(cell, mapa: dict, *, cor_padrao=COR_CINZA,
                                       size_padrao=10):
    """
    Se a célula contém exatamente um placeholder do `mapa`,
    substitui pelo valor com cor/estilo apropriado.
    Suporta vários placeholders na mesma célula (linha de assinatura, ex.).
    """
    texto_completo = cell.text
    placeholders_encontrados = [(ph, val, cor, sz) for ph, (val, cor, sz) in mapa.items()
                                 if f"[{ph}]" in texto_completo]
    if not placeholders_encontrados:
        return

    # Caso especial: célula só contém um placeholder isolado
    if len(placeholders_encontrados) == 1:
        ph, val, cor, sz = placeholders_encontrados[0]
        bare = texto_completo.strip() == f"[{ph}]"
        if bare:
            _set_celula(cell, val if val else PLACEHOLDER_VAZIO,
                        cor=cor, bold=True, size=sz)
            return

    # Caso geral: várias substituições no texto (mantém estrutura original)
    novo_texto = texto_completo
    for ph, val, _cor, _sz in placeholders_encontrados:
        novo_texto = novo_texto.replace(f"[{ph}]", val if val else PLACEHOLDER_VAZIO)
    # Preserva quebras de linha do template
    paragraphs = cell.paragraphs
    if len(paragraphs) == 1:
        # Só um parágrafo — refaz com o novo texto
        p = paragraphs[0]
        # Tenta manter alinhamento atual
        antigo_align = p.alignment
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        # Quebra em linhas, mantendo as quebras como linhas separadas
        linhas = novo_texto.split('\n')
        for i, linha in enumerate(linhas):
            run = p.add_run(linha)
            run.font.color.rgb = cor_padrao
            run.font.size = Pt(size_padrao)
            if i < len(linhas) - 1:
                run.add_break()
        if antigo_align is not None:
            p.alignment = antigo_align
    else:
        # Múltiplos parágrafos — substitui texto de cada um sem refazer estrutura
        partes = novo_texto.split('\n')
        for i, p in enumerate(paragraphs):
            txt = partes[i] if i < len(partes) else ""
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            run = p.add_run(txt)
            run.font.color.rgb = cor_padrao
            run.font.size = Pt(size_padrao)


def _substituir_em_paragrafo(par, mapa: dict, *, cor=COR_CINZA, size=10):
    if not par.text:
        return
    novo = par.text
    mudou = False
    for ph, (val, _cor, _sz) in mapa.items():
        token = f"[{ph}]"
        if token in novo:
            novo = novo.replace(token, val if val else PLACEHOLDER_VAZIO)
            mudou = True
    if mudou:
        for r in list(par.runs):
            r._element.getparent().remove(r._element)
        run = par.add_run(novo)
        run.font.color.rgb = cor
        run.font.size = Pt(size)


# ─────────────────────────────────────────────────────────────────────────────
# Inserção de fotos
# ─────────────────────────────────────────────────────────────────────────────
def _inserir_fotos_no_marcador(doc: Document, marcador: str, fotos: list[str],
                                largura_cm: float = 7.5):
    """Procura um parágrafo cujo texto seja `marcador` e o substitui por imagens."""
    fotos = [f for f in (fotos or []) if f and Path(f).exists()]
    alvo = None
    for p in doc.paragraphs:
        if p.text.strip() == marcador:
            alvo = p
            break
    if alvo is None:
        return

    # Limpa o texto do parágrafo
    for r in list(alvo.runs):
        r._element.getparent().remove(r._element)

    if not fotos:
        run = alvo.add_run("(sem fotos registradas)")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = COR_CINZA_CLARO
        return

    # Insere a primeira foto no parágrafo existente
    alvo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = alvo.add_run()
    try:
        run.add_picture(fotos[0], width=Cm(largura_cm))
    except Exception as e:
        run.add_text(f"(erro ao inserir {Path(fotos[0]).name}: {e})")

    # Insere as restantes em parágrafos novos, logo depois
    parent = alvo._element.getparent()
    idx = list(parent).index(alvo._element)
    for foto in fotos[1:]:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        novo_p = OxmlElement('w:p')
        parent.insert(idx + 1, novo_p)
        idx += 1
        # Cria parágrafo via API e move para a posição
        # (caminho mais simples: usa add_paragraph e depois reordena)
    # Estratégia mais simples — usa doc.add_paragraph + reposicionar
    # Refazendo: usar uma abordagem mais simples
    if len(fotos) > 1:
        # Remove os elementos vazios que adicionamos acima
        for child in list(parent)[idx - (len(fotos) - 1) + 1 : idx + 1]:
            parent.remove(child)
        # Agora adiciona via doc.add_paragraph e move
        depois = []
        for foto in fotos[1:]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.add_run().add_picture(foto, width=Cm(largura_cm))
            except Exception:
                p.add_run(f"(erro ao inserir {Path(foto).name})")
            depois.append(p._element)
            parent.remove(p._element)
        # Insere depois do parágrafo alvo
        ref = alvo._element
        for novo in depois:
            ref.addnext(novo)
            ref = novo


# ─────────────────────────────────────────────────────────────────────────────
# Função principal
# ─────────────────────────────────────────────────────────────────────────────
def gerar_vistoria_entrada_saida(
    dados: dict,
    caminho_saida: str,
    template_path: Optional[str] = None,
) -> dict:
    """
    Gera o DOCX a partir do template e devolve um resumo.
    """
    template = Path(template_path or
                    Path(__file__).parent / "docx_templates" /
                    "VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx")
    if not template.exists():
        raise FileNotFoundError(f"Template não encontrado: {template}")

    acc_in  = dados.get("acessorios_entrada") or {}
    acc_out = dados.get("acessorios_saida")   or {}

    tem_saida = any([
        dados.get("data_saida"), dados.get("hodometro_saida"),
        dados.get("combustivel_saida"), acc_out,
        dados.get("obs_saida"), dados.get("sintomas_saida"),
        dados.get("fotos_saida"),
    ])
    status_doc = "completa" if tem_saida else "pendente_saida"

    # ── Monta o mapa de substituições ─────────────────────────────────────
    # mapa[placeholder] = (valor, cor, tamanho)
    mapa: dict[str, tuple[str, RGBColor, int]] = {}

    campos_simples = [
        "contrato_id", "cliente_nome", "cliente_telefone",
        "cliente_endereco", "preenchido_por",
        "veiculo", "placa", "cor", "ano", "chassi", "numero_motor",
        "data_entrada", "data_saida",
        "hodometro_entrada", "hodometro_saida",
        "combustivel_entrada", "combustivel_saida",
        "obs_entrada", "obs_saida",
        "sintomas_entrada", "sintomas_saida",
        "responsavel_entrada", "responsavel_saida",
    ]
    for campo in campos_simples:
        mapa[campo] = (str(dados.get(campo) or ""), COR_CINZA, 10)

    # Acessórios: cada um vira 3 placeholders
    divergencias = []
    for key in ACESSORIOS:
        e = acc_in.get(key, "")
        s = acc_out.get(key, "")
        status_txt, status_cor = _calcular_status_acessorio(e, s)
        mapa[f"{key}_entrada"] = (e, _cor_valor(e), 9)
        mapa[f"{key}_saida"]   = (s, _cor_valor(s), 9)
        mapa[f"{key}_status"]  = (status_txt, status_cor, 9)

        if status_txt and "OK" not in status_txt and "Aguardando" not in status_txt:
            label = key.replace("acc_", "").replace("_", " ").title()
            divergencias.append((label, e or "—", s or "—", status_txt))

    # ── Abre o template e aplica substituições ────────────────────────────
    doc = Document(str(template))

    # Em todas as tabelas
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                _substituir_placeholder_no_celula(cell, mapa)

    # Em todos os parágrafos do corpo (cabeçalho, subtítulo, "Contrato X")
    for par in doc.paragraphs:
        _substituir_em_paragrafo(par, mapa)

    # ── Inserir fotos ─────────────────────────────────────────────────────
    _inserir_fotos_no_marcador(doc, "[FOTOS_ENTRADA]",
                                dados.get("fotos_entrada", []))
    _inserir_fotos_no_marcador(doc, "[FOTOS_SAIDA]",
                                dados.get("fotos_saida", []))

    # ── Salvar ────────────────────────────────────────────────────────────
    saida = Path(caminho_saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(saida))

    return {
        "arquivo":      str(saida),
        "status":       status_doc,
        "divergencias": divergencias,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Conversão para PDF (mesma abordagem do gerar_contrato.py)
# ─────────────────────────────────────────────────────────────────────────────
def docx_para_pdf(caminho_docx: str) -> Optional[str]:
    """
    Converte o DOCX para PDF usando docx2pdf (Windows + Word) ou
    libreoffice (Linux). Retorna o caminho do PDF gerado, ou None.
    """
    import platform, subprocess
    caminho_pdf = str(Path(caminho_docx).with_suffix(".pdf"))
    sistema = platform.system().lower()
    try:
        if sistema == "windows":
            from docx2pdf import convert
            convert(caminho_docx, caminho_pdf)
        else:
            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", str(Path(caminho_docx).parent),
                caminho_docx
            ], check=True, timeout=60)
        return caminho_pdf if Path(caminho_pdf).exists() else None
    except Exception as e:
        print(f"[gerar_vistoria_entrada_saida] PDF falhou: {e}")
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Execução direta — gera um exemplo
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    exemplo = {
        "contrato_id":       "0042",
        "cliente_nome":      "JOÃO DA SILVA SANTOS",
        "cliente_telefone":  "(83) 99999-0000",
        "cliente_endereco":  "Rua das Acácias, 123 — Bessa, João Pessoa/PB",
        "preenchido_por":    "Ana Karolina",
        "veiculo":           "VW Gol Special 1.0",
        "placa":             "ELY-4D83",
        "cor":               "BRANCO",
        "ano":               "2022/2023",
        "chassi":            "9BWZZZ377VT004251",
        "numero_motor":      "CHZ123456",

        "data_entrada":         "10/05/2026 14:30",
        "hodometro_entrada":    "32.450",
        "combustivel_entrada":  "3/4",
        "obs_entrada":          "Veículo limpo, sem riscos visíveis.",
        "sintomas_entrada":     "",
        "responsavel_entrada":  "Ana Karolina",
        "acessorios_entrada":   {k: "S" for k in ACESSORIOS},

        "data_saida":         "20/05/2026 11:00",
        "hodometro_saida":    "33.120",
        "combustivel_saida":  "1/4",
        "obs_saida":          "Para-choque dianteiro arranhado no canto esquerdo.",
        "sintomas_saida":     "Ar-condicionado fazendo barulho ao ligar.",
        "responsavel_saida":  "Ana Karolina",
        "acessorios_saida": {
            **{k: "S" for k in ACESSORIOS},
            "acc_calotas":   "N",   # perdeu uma calota
            "acc_tapetes":   "A",   # tapetes avariados
            "acc_estepe":    "N",   # estepe sumiu
            "acc_triangulo": "N",
        },
        "fotos_entrada": [],
        "fotos_saida":   [],
    }
    saida = sys.argv[1] if len(sys.argv) > 1 else "VISTORIA_EXEMPLO.docx"
    resumo = gerar_vistoria_entrada_saida(
        exemplo,
        caminho_saida=saida,
        template_path="VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx",
    )
    print(f"Gerado: {resumo['arquivo']}")
    print(f"Status: {resumo['status']}")
    print(f"Divergências encontradas: {len(resumo['divergencias'])}")
    for label, e, s, motivo in resumo['divergencias']:
        print(f"  • {label}: {e} → {s}  ({motivo})")
